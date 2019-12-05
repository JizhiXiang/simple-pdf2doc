# -*- coding: utf-8 -*- 

import os
from configparser import ConfigParser
from io import StringIO
from io import open

from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfinterp import process_pdf
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from docx import Document

def read_from_pdf(file_path):
    with open(file_path, 'rb') as file:
        resource_manager = PDFResourceManager()
        return_str = StringIO()
        lap_params = LAParams()

        device = TextConverter(
            resource_manager, return_str, laparams=lap_params)
        process_pdf(resource_manager, device, file)
        device.close()
        content = return_str.getvalue()
        return_str.close()
        return content

def save_text_to_word(content, file_path):
    doc = Document()
    for line in content.split('\n'):
        paragraph = doc.add_paragraph()
        paragraph.add_run(remove_control_characters(line))
    doc.save(file_path)

def remove_control_characters(content):
    mpa = dict.fromkeys(range(32))
    return content.translate(mpa)

def pdf_to_word(pdf_file_path, word_file_path):
    content = read_from_pdf(pdf_file_path)
    save_text_to_word(content, word_file_path)

def main(pdf, word="out.docx"):
    try:
        pdf_to_word(pdf, word)
        print("转换成功！")
    except:
        print("出错了")

if __name__ == '__main__':
    # pdf = input()
    pdf = "E:/aaa/bbb/your_file.pdf"
    main(pdf)